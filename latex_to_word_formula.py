#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
latex_to_word_formula.py  v2.0
================================
批量将 Word 文档中所有 LaTeX 公式转换为 Word 原生可编辑公式（OMML 格式）。

核心改进（v2.0）：
  - 修复 Word 内部将 \\ 双写存储导致正则无法匹配的问题
  - 支持公式跨多个 run 分布的情况（段落级扫描）
  - 段落 XML 整体重建策略，彻底避免 run 边界问题

支持的公式定界符：
  行内公式：  $...$   \\(...\\)
  块级公式：  $$...$$  \\[...\\]

使用方法：
  python latex_to_word_formula.py input.docx
  python latex_to_word_formula.py input.docx -o output.docx
  python latex_to_word_formula.py input.docx --no-xslt

依赖：
  pip install python-docx lxml latex2mathml
"""

import re
import sys
import copy
import argparse
from pathlib import Path
from lxml import etree
import latex2mathml.converter
from docx import Document
from docx.oxml.ns import qn

# ──────────────────────────────────────────────
#  命名空间常量
# ──────────────────────────────────────────────
W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M  = "http://schemas.openxmlformats.org/officeDocument/2006/math"
ML = "http://www.w3.org/1998/Math/MathML"

def wtag(name): return f"{{{W}}}{name}"
def mtag(name): return f"{{{M}}}{name}"

# ──────────────────────────────────────────────
#  XSLT 加载（Office 官方 MML2OMML）
# ──────────────────────────────────────────────
def _load_mml2omml_xslt():
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office15\MML2OMML.XSL",
    ]
    for p in candidates:
        if Path(p).exists():
            print(f"  [XSLT] Office 官方转换器: {p}")
            with open(p, "rb") as f:
                return etree.XSLT(etree.parse(f))
    print("  [XSLT] 未找到 Office 安装，使用内置转换器")
    return None


def _latex_to_mml(latex_str: str) -> str:
    """LaTeX -> MathML 字符串，处理常见异常。"""
    # latex2mathml 不支持 \left \right 的部分写法，预处理
    tex = latex_str
    try:
        return latex2mathml.converter.convert(tex)
    except Exception:
        pass
    # 简单清理后重试
    tex2 = re.sub(r'\\left\s*\.', r'\\left(', tex)
    tex2 = re.sub(r'\\right\s*\.', r'\\right)', tex2)
    try:
        return latex2mathml.converter.convert(tex2)
    except Exception as e:
        return f'<math xmlns="{ML}"><mtext>{_escape_xml(latex_str)}</mtext></math>'


def _escape_xml(s):
    return s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')


def make_omml(latex_str: str, is_display: bool, xslt=None) -> etree._Element:
    """LaTeX -> OMML Element (m:oMath 或 m:oMathPara)。"""
    mml_str = _latex_to_mml(latex_str)
    if xslt is not None:
        try:
            mml_doc = etree.fromstring(mml_str.encode("utf-8"))
            omml_doc = xslt(mml_doc)
            root = omml_doc.getroot()
            if is_display:
                # 包在 oMathPara 里
                tag = etree.QName(root).localname
                if tag != "oMathPara":
                    para = etree.Element(mtag("oMathPara"))
                    para.append(root)
                    return para
            return root
        except Exception as e:
            print(f"    [WARN] XSLT 失败({e})，回退内置转换")
    return _builtin_convert(mml_str, is_display)


# ──────────────────────────────────────────────
#  内置 MathML -> OMML 转换
# ──────────────────────────────────────────────
def _builtin_convert(mml_str: str, is_display: bool) -> etree._Element:
    try:
        root = etree.fromstring(mml_str.encode("utf-8"))
    except Exception:
        oMath = etree.Element(mtag("oMath"))
        r = etree.SubElement(oMath, mtag("r"))
        t = etree.SubElement(r, mtag("t"))
        t.text = mml_str
        return oMath
    oMath = etree.Element(mtag("oMath"))
    for el in _mml2omml(root):
        oMath.append(el)
    if is_display:
        para = etree.Element(mtag("oMathPara"))
        para.append(oMath)
        return para
    return oMath


def _mml2omml(el):
    """递归将 MathML Element 转换为 OMML Element 列表。"""
    if el.tag == etree.Comment:
        return []
    tag = etree.QName(el).localname
    ch  = list(el)

    def kids():
        out = []
        for c in ch: out.extend(_mml2omml(c))
        return out

    def txt():
        s = (el.text or "").strip()
        if not s: return []
        r = etree.Element(mtag("r"))
        if tag == "mi" and len(s) == 1:
            rPr = etree.SubElement(r, mtag("rPr"))
            sty = etree.SubElement(rPr, mtag("sty"))
            sty.set(mtag("val"), "i")
        t = etree.SubElement(r, mtag("t"))
        t.text = s
        return [r]

    if tag in ("math","mrow","mstyle","mpadded","mphantom","merror","menclose","maction"):
        return kids()

    if tag in ("mi","mn","mo","mtext","ms"):
        return txt()

    if tag == "msup" and len(ch) >= 2:
        e = etree.Element(mtag("sSup"))
        eb = etree.SubElement(e, mtag("e"))
        sb = etree.SubElement(e, mtag("sup"))
        for x in _mml2omml(ch[0]): eb.append(x)
        for x in _mml2omml(ch[1]): sb.append(x)
        return [e]

    if tag == "msub" and len(ch) >= 2:
        e = etree.Element(mtag("sSub"))
        eb = etree.SubElement(e, mtag("e"))
        sb = etree.SubElement(e, mtag("sub"))
        for x in _mml2omml(ch[0]): eb.append(x)
        for x in _mml2omml(ch[1]): sb.append(x)
        return [e]

    if tag == "msubsup" and len(ch) >= 3:
        e = etree.Element(mtag("sSubSup"))
        eb  = etree.SubElement(e, mtag("e"))
        sb  = etree.SubElement(e, mtag("sub"))
        spb = etree.SubElement(e, mtag("sup"))
        for x in _mml2omml(ch[0]): eb.append(x)
        for x in _mml2omml(ch[1]): sb.append(x)
        for x in _mml2omml(ch[2]): spb.append(x)
        return [e]

    if tag == "mfrac" and len(ch) >= 2:
        f = etree.Element(mtag("f"))
        n = etree.SubElement(f, mtag("num"))
        d = etree.SubElement(f, mtag("den"))
        for x in _mml2omml(ch[0]): n.append(x)
        for x in _mml2omml(ch[1]): d.append(x)
        return [f]

    if tag == "msqrt":
        rad = etree.Element(mtag("rad"))
        pr  = etree.SubElement(rad, mtag("radPr"))
        dh  = etree.SubElement(pr,  mtag("degHide"))
        dh.set(mtag("val"), "1")
        etree.SubElement(rad, mtag("deg"))
        e = etree.SubElement(rad, mtag("e"))
        for x in kids(): e.append(x)
        return [rad]

    if tag == "mroot" and len(ch) >= 2:
        rad = etree.Element(mtag("rad"))
        d   = etree.SubElement(rad, mtag("deg"))
        e   = etree.SubElement(rad, mtag("e"))
        for x in _mml2omml(ch[1]): d.append(x)
        for x in _mml2omml(ch[0]): e.append(x)
        return [rad]

    if tag == "mover" and len(ch) >= 2:
        lu = etree.Element(mtag("limUpp"))
        e  = etree.SubElement(lu, mtag("e"))
        l  = etree.SubElement(lu, mtag("lim"))
        for x in _mml2omml(ch[0]): e.append(x)
        for x in _mml2omml(ch[1]): l.append(x)
        return [lu]

    if tag == "munder" and len(ch) >= 2:
        ll = etree.Element(mtag("limLow"))
        e  = etree.SubElement(ll, mtag("e"))
        l  = etree.SubElement(ll, mtag("lim"))
        for x in _mml2omml(ch[0]): e.append(x)
        for x in _mml2omml(ch[1]): l.append(x)
        return [ll]

    if tag == "munderover" and len(ch) >= 3:
        ny = etree.Element(mtag("nary"))
        pr = etree.SubElement(ny, mtag("naryPr"))
        ll = etree.SubElement(pr, mtag("limLoc"))
        ll.set(mtag("val"), "undOvr")
        sb  = etree.SubElement(ny, mtag("sub"))
        spb = etree.SubElement(ny, mtag("sup"))
        e   = etree.SubElement(ny, mtag("e"))
        for x in _mml2omml(ch[1]): sb.append(x)
        for x in _mml2omml(ch[2]): spb.append(x)
        return [ny]

    if tag == "mtable":
        m = etree.Element(mtag("m"))
        etree.SubElement(etree.SubElement(m, mtag("mPr")), mtag("mcs"))
        for row in ch:
            if etree.QName(row).localname == "mtr":
                mr = etree.SubElement(m, mtag("mr"))
                for cell in row:
                    if etree.QName(cell).localname == "mtd":
                        me = etree.SubElement(mr, mtag("e"))
                        for x in _mml2omml(cell): me.append(x)
        return [m]

    if tag == "mfenced":
        ob = el.get("open","("); cb = el.get("close",")")
        d  = etree.Element(mtag("d"))
        pr = etree.SubElement(d, mtag("dPr"))
        bc = etree.SubElement(pr, mtag("begChr")); bc.set(mtag("val"), ob)
        ec = etree.SubElement(pr, mtag("endChr")); ec.set(mtag("val"), cb)
        for c in ch:
            e = etree.SubElement(d, mtag("e"))
            for x in _mml2omml(c): e.append(x)
        return [d]

    return txt() or kids()


# ──────────────────────────────────────────────
#  公式正则（同时匹配单/双反斜杠存储形式）
# ──────────────────────────────────────────────
# Word 在 .docx 内部有时会将 \ 存为 \\（双写），
# 因此定界符 \( \) \[ \] 可能呈现为 \\( \\) \\[ \\]
# 我们两种形式都要匹配。

_RE = re.compile(
    # ① 块级 $$...$$
    r'\$\$(?P<d1>[\s\S]+?)\$\$'
    # ② 块级 \[...\] 或 \\[...\\]
    r'|(?:\\{1,2})\[(?P<d2>[\s\S]+?)(?:\\{1,2})\]'
    # ③ 行内 \(...\) 或 \\(...\\)
    r'|(?:\\{1,2})\((?P<i1>[\s\S]+?)(?:\\{1,2})\)'
    # ④ 行内 $...$ （非 $$）
    r'|(?<!\$)\$(?!\$)(?P<i2>[^\n$]+?)(?<!\$)\$(?!\$)'
)


def _find(text):
    """返回 [(start, end, latex, is_display), ...]"""
    out = []
    for m in _RE.finditer(text):
        if m.group("d1") is not None:
            out.append((m.start(), m.end(), m.group("d1").strip(), True))
        elif m.group("d2") is not None:
            out.append((m.start(), m.end(), m.group("d2").strip(), True))
        elif m.group("i1") is not None:
            out.append((m.start(), m.end(), m.group("i1").strip(), False))
        elif m.group("i2") is not None:
            out.append((m.start(), m.end(), m.group("i2").strip(), False))
    return out


def _normalize_latex(s: str) -> str:
    """
    将 Word 存储的双反斜杠还原为单反斜杠，
    例如 \\\\frac -> \\frac，\\\\alpha -> \\alpha。
    """
    # 仅当字符串中出现连续4个反斜杠（即存储层的\\）时才替换
    # 策略：把 \\\\ 替换为 \\（两个 Python 反斜杠，即实际一个）
    # 注意：python 字符串里 '\\\\' 代表两个实际反斜杠
    result = s.replace('\\\\', '\\')
    return result


# ──────────────────────────────────────────────
#  段落级处理：合并所有 run 文字 → 识别公式 → 重建 XML
# ──────────────────────────────────────────────

def _get_para_runs(para_xml):
    """
    返回段落直属子 run 列表（不含嵌套在 hyperlink 等里面的），
    每项: (run_element, text_or_None)
    """
    result = []
    for child in para_xml:
        local = etree.QName(child).localname if child.tag != etree.Comment else None
        if local == "r":
            t = child.find(wtag("t"))
            result.append((child, t.text if t is not None else None))
        else:
            result.append((child, None))   # 非 run 子节点（如 pPr、bookmarkStart 等）
    return result


def _rebuild_para(para_xml, xslt, counter):
    """
    对段落进行整体重建：
    1. 收集所有 run 的文字（保留各 run 的 rPr）
    2. 合并为连续字符串，扫描公式
    3. 按公式位置切割，重新生成 run 和 OMML 节点
    4. 替换原段落中的 run 序列
    返回本段转换的公式数。
    """
    # ── Step 1：找出所有直属子节点 ──
    children = list(para_xml)
    # 收集"run 组"的索引范围和文字
    # 只处理直属 w:r（忽略 hyperlink 等复杂嵌套）
    
    run_info = []   # (index_in_children, run_el, text)
    for idx, child in enumerate(children):
        if child.tag == wtag("r"):
            t_el = child.find(wtag("t"))
            txt  = t_el.text if t_el is not None else ""
            run_info.append((idx, child, txt or ""))

    if not run_info:
        return 0

    # ── Step 2：合并文字，同时记录每个字符属于哪个 run ──
    combined = ""
    char_map = []   # char_map[i] = run_index in run_info
    for ri, (_, _, txt) in enumerate(run_info):
        for _ in txt:
            char_map.append(ri)
        combined += txt

    # ── Step 3：在合并文字中扫描公式 ──
    matches = _find(combined)
    if not matches:
        return 0

    # ── Step 4：对每个公式段，检查是否跨 run ──
    # 先将找到的公式 latex 还原反斜杠
    fixed = []
    for (s, e, latex, is_disp) in matches:
        fixed.append((s, e, _normalize_latex(latex), is_disp))
    matches = fixed

    # ── Step 5：重建节点序列 ──
    # 策略：从 combined 中按字符位置切割，
    # 文字段 → 使用首个覆盖 run 的 rPr 重建 w:r
    # 公式段 → 生成 OMML 节点

    def _make_text_run(text, sample_run):
        """创建文字 run，复制 sample_run 的 rPr。"""
        if not text:
            return None
        nr = etree.Element(wtag("r"))
        rPr = sample_run.find(wtag("rPr"))
        if rPr is not None:
            nr.append(copy.deepcopy(rPr))
        t = etree.SubElement(nr, wtag("t"))
        t.text = text
        # 保留首尾空格
        if text != text.strip():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        return nr

    def _run_at(char_idx):
        """返回 combined[char_idx] 对应的 run element。"""
        if char_idx >= len(char_map):
            return run_info[-1][1]
        ri = char_map[char_idx]
        return run_info[ri][1]

    new_nodes = []   # 最终替换 run 区间的新节点列表
    cursor = 0

    for (start, end, latex, is_disp) in matches:
        # 公式前的文字
        if cursor < start:
            seg = combined[cursor:start]
            sample = _run_at(cursor)
            nr = _make_text_run(seg, sample)
            if nr is not None:
                new_nodes.append(nr)
        # 公式本身
        print(f"    [转换] {'块级' if is_disp else '行内'}: {latex[:70]}{'...' if len(latex)>70 else ''}")
        counter[0] += 1
        try:
            omml_el  = make_omml(latex, is_disp, xslt)
            omml_xml = etree.fromstring(etree.tostring(omml_el, encoding="unicode").encode("utf-8"))
            new_nodes.append(omml_xml)
        except Exception as ex:
            print(f"    [ERROR] 转换失败: {ex}，保留原文")
            sample = _run_at(start)
            nr = _make_text_run(combined[start:end], sample)
            if nr is not None:
                new_nodes.append(nr)
            counter[0] -= 1
        cursor = end

    # 公式后剩余文字
    if cursor < len(combined):
        seg = combined[cursor:]
        sample = _run_at(cursor if cursor < len(char_map) else len(char_map)-1)
        nr = _make_text_run(seg, sample)
        if nr is not None:
            new_nodes.append(nr)

    # ── Step 6：找出 run 区间在 para_xml 中的位置并替换 ──
    # 找到第一个和最后一个 run 的位置
    first_run_idx = run_info[0][0]
    last_run_idx  = run_info[-1][0]

    # 删除原有的所有直属 run
    for (idx, run_el, _) in run_info:
        para_xml.remove(run_el)

    # 在 first_run_idx 处插入新节点
    # （删除后索引会偏移，重新找插入位置）
    # 非 run 节点保持原位，我们把新节点插在第一个 run 原来的位置
    # 用更安全的方式：找到第一个非 run 节点之后插入
    insert_after = None
    for child in para_xml:
        lc = etree.QName(child).localname if child.tag != etree.Comment else ""
        if lc in ("pPr", "bookmarkStart", "bookmarkEnd", "proofErr"):
            insert_after = child
        else:
            break

    if insert_after is not None:
        insert_pos = list(para_xml).index(insert_after) + 1
    else:
        insert_pos = 0

    for i, node in enumerate(new_nodes):
        para_xml.insert(insert_pos + i, node)

    return len(matches)


# ──────────────────────────────────────────────
#  文档级处理
# ──────────────────────────────────────────────

def process_doc(doc: Document, xslt) -> int:
    counter = [0]

    def _do_para(para_xml):
        # 快速检查：段落文字里有公式定界符吗？
        all_t = "".join(
            (t.text or "") for t in para_xml.findall(".//" + wtag("t"))
        )
        if not any(c in all_t for c in ("$", r"\("[:2], r"\["[:2])):
            return
        n = _rebuild_para(para_xml, xslt, counter)
        if n:
            print(f"    -> 本段转换 {n} 个公式")

    print("\n[处理正文段落]")
    for i, para in enumerate(doc.paragraphs):
        _do_para(para._element)

    print("\n[处理表格]")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _do_para(para._element)

    print("\n[处理页眉页脚]")
    for section in doc.sections:
        for hf_obj in [section.header, section.footer,
                       section.even_page_header, section.even_page_footer,
                       section.first_page_header, section.first_page_footer]:
            try:
                for para in hf_obj.paragraphs:
                    _do_para(para._element)
            except Exception:
                pass

    return counter[0]


# ──────────────────────────────────────────────
#  主程序
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="批量将 Word 文档中的 LaTeX 公式转换为 Word 原生 OMML 公式",
    )
    parser.add_argument("input",  help="输入 Word 文档路径 (.docx)")
    parser.add_argument("-o", "--output", help="输出文档路径（默认：原名_converted.docx）")
    parser.add_argument("--no-xslt", action="store_true", help="强制使用内置转换器")
    args = parser.parse_args()

    inp = Path(args.input)
    if not inp.exists():
        print(f"[错误] 文件不存在: {inp}"); sys.exit(1)
    if inp.suffix.lower() != ".docx":
        print(f"[错误] 仅支持 .docx，当前: {inp.suffix}"); sys.exit(1)

    out = Path(args.output) if args.output else inp.with_name(inp.stem + "_converted" + inp.suffix)

    print("=" * 60)
    print(f"  输入: {inp}")
    print(f"  输出: {out}")
    print("=" * 60)

    xslt = None if args.no_xslt else _load_mml2omml_xslt()

    print(f"\n[加载文档] {inp.name}")
    doc = Document(str(inp))

    total = process_doc(doc, xslt)

    print(f"\n[保存] -> {out.name}")
    doc.save(str(out))

    print("\n" + "=" * 60)
    print(f"  [完成] 共转换 {total} 个公式")
    print(f"  输出: {out.resolve()}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
