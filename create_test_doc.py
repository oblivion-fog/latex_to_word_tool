#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""创建包含 LaTeX 公式的测试 Word 文档"""

from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading("LaTeX 公式转换测试文档", 0)

doc.add_paragraph("本文档用于测试 LaTeX 公式自动转换功能。")

# 行内公式
p = doc.add_paragraph()
p.add_run("行内公式示例：设 ")
p.add_run("$x^2 + y^2 = r^2$")
p.add_run(" 为圆的方程，其中 ")
p.add_run(r"\(E = mc^2\)")
p.add_run(" 是爱因斯坦质能方程。")

# 分式
p = doc.add_paragraph()
p.add_run("分式示例：$\\frac{a+b}{c-d}$ 以及 $\\frac{\\partial f}{\\partial x}$")

# 根号
p = doc.add_paragraph()
p.add_run("根号示例：$\\sqrt{x^2+y^2}$ 和 $\\sqrt[3]{a+b}$")

# 上下标
p = doc.add_paragraph()
p.add_run("上下标：$x_i^2$ 和 $\\sum_{i=1}^{n} x_i$")

# 块级公式
p = doc.add_paragraph()
p.add_run("块级公式（积分）：")
p.add_run("$$\\int_{-\\infty}^{+\\infty} e^{-x^2} dx = \\sqrt{\\pi}$$")

# 希腊字母
p = doc.add_paragraph()
p.add_run("希腊字母：$\\alpha + \\beta = \\gamma$，$\\Delta x \\to 0$")

# 矩阵
p = doc.add_paragraph()
p.add_run(r"矩阵示例：$$\begin{pmatrix} a & b \\ c & d \end{pmatrix}$$")

# 复杂公式
p = doc.add_paragraph()
p.add_run(r"复杂示例：$$\lim_{n\to\infty} \left(1 + \frac{1}{n}\right)^n = e$$")

doc.save(r"c:\Users\miwus\Desktop\Math Model\test_latex.docx")
print("测试文档已创建：test_latex.docx")
